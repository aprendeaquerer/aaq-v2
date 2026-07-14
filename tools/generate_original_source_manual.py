#!/usr/bin/env python3
import csv
import json
import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    ListFlowable,
    ListItem,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE_ROOT = Path("/Users/pedro/Desktop/Archivo")
OUT = ROOT / "output"
PDF_OUT = OUT / "pdf"
TRACE_OUT = OUT / "traceability"
SOURCE_OUT = OUT / "original_sources"
TEXT_OUT = SOURCE_OUT / "extracted_text"


CHAPTERS = [
    {
        "title": "Tesis editorial y arquitectura del manual definitivo",
        "keywords": ["attachment", "relationship", "inner work", "emotion", "boundaries", "polarity"],
        "focus": "criterios de integracion, profundidad, trazabilidad y conservacion de matices",
    },
    {
        "title": "Fundamentos del apego adulto",
        "keywords": ["attachment", "secure", "anxious", "avoidant", "ainsworth", "levine", "origin"],
        "focus": "base teorica, estilos, seguridad, protesta, evitacion y sistema de vinculo",
    },
    {
        "title": "Origen psicologico de los patrones relacionales",
        "keywords": ["wound", "childhood", "needs", "feelings", "trauma", "injury", "rejection"],
        "focus": "heridas, experiencias tempranas, necesidades no expresadas y aprendizaje relacional",
    },
    {
        "title": "Mecanismos emocionales, corporales y de regulacion",
        "keywords": ["emotion regulation", "nervous", "body", "regulation", "somatic", "emotion"],
        "focus": "regulacion emocional, activacion fisiologica, cuerpo y tolerancia al malestar",
    },
    {
        "title": "Ciclo negativo, conflicto y reparacion",
        "keywords": ["negative cycle", "repair", "conflict", "secure love", "responding", "shame"],
        "focus": "persecucion-retirada, protesta, defensividad, verguenza, reaching y responding",
    },
    {
        "title": "Evitacion, desorganizacion y miedo a la intimidad",
        "keywords": ["avoidant", "disorganized", "therapy", "falling out", "falls in love", "intimacy"],
        "focus": "estrategias desactivadoras, ambivalencia, desconexion y reconstruccion de seguridad",
    },
    {
        "title": "Ansiedad, rechazo e inversion emocional prematura",
        "keywords": ["rejected", "emotionally invested", "mindset", "breakup", "potential"],
        "focus": "apego ansioso, sobreinversion, rechazo, ruptura, duelo y eleccion de pareja",
    },
    {
        "title": "Comunicacion, limites, compasion y necesidades",
        "keywords": ["boundaries", "compassion", "ask", "needs", "feelings", "know", "relationship"],
        "focus": "pedidos claros, limites, compasion sin autoabandono y conversaciones reparadoras",
    },
    {
        "title": "Discernimiento, dating y eleccion de pareja",
        "keywords": ["dating", "settling", "potential", "find love", "move on", "ex"],
        "focus": "senales, estandares, compatibilidad, disponibilidad y salir de vinculos no nutritivos",
    },
    {
        "title": "Inner work, identidad y agencia personal",
        "keywords": ["inner work", "identity", "change", "self", "relationships", "wounds"],
        "focus": "trabajo interno, identidad, agencia, valores, autoobservacion y practica sostenida",
    },
    {
        "title": "Polaridad, atraccion y dinamicas masculino/femenino",
        "keywords": ["masculine", "feminine", "polarity", "attraction", "leader"],
        "focus": "polaridad no dogmatica, atraccion, liderazgo emocional, deseo y reciprocidad",
    },
    {
        "title": "Consejos y principios para hombres / Masculino",
        "keywords": ["masculine", "man", "leader", "connor beaton", "male"],
        "focus": "material dirigido a hombres o energia masculina, separado del cuerpo general",
    },
    {
        "title": "Herramientas y ejercicios no meditativos",
        "keywords": ["steps", "practice", "exercise", "template", "process", "routine"],
        "focus": "ejercicios de reflexion, escritura, comunicacion, limites y toma de decisiones; sin meditaciones ni breathwork",
    },
]


def clean(value):
    value = re.sub(r"\s+", " ", value or "").strip()
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def safe_name(path):
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", path.relative_to(SOURCE_ROOT).as_posix())


def extract_pdf(path):
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            text = f"[Error extrayendo pagina {i}: {exc}]"
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages), len(reader.pages)


def extract_docx(path):
    doc = Document(str(path))
    blocks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n\n".join(blocks), None


def iter_sources():
    for path in sorted(SOURCE_ROOT.rglob("*")):
        if path.name.startswith(".") or not path.is_file():
            continue
        if path.suffix.lower() in {".pdf", ".docx"}:
            yield path


def classify(path, text):
    haystack = f"{path.as_posix()} {text[:20000]}".lower()
    scores = {}
    for chapter in CHAPTERS:
        score = sum(haystack.count(k.lower()) for k in chapter["keywords"])
        scores[chapter["title"]] = score
    primary = max(scores, key=scores.get)
    if scores[primary] == 0:
        primary = "Tesis editorial y arquitectura del manual definitivo"
    return primary, scores


def extract_sources():
    TEXT_OUT.mkdir(parents=True, exist_ok=True)
    sources = []
    for path in iter_sources():
        if path.suffix.lower() == ".pdf":
            text, pages = extract_pdf(path)
        else:
            text, pages = extract_docx(path)
        text_path = TEXT_OUT / f"{safe_name(path)}.txt"
        text_path.write_text(text, encoding="utf-8")
        primary, scores = classify(path, text)
        sources.append({
            "path": str(path),
            "relative_path": path.relative_to(SOURCE_ROOT).as_posix(),
            "title": path.stem,
            "type": path.suffix.lower().lstrip("."),
            "pages": pages,
            "word_count": len(re.findall(r"\w+", text)),
            "char_count": len(text),
            "text_path": str(text_path),
            "primary_chapter": primary,
            "chapter_scores": scores,
        })
    return sources


def chapter_sources(sources, chapter):
    title = chapter["title"]
    ranked = []
    for source in sources:
        score = source["chapter_scores"].get(title, 0)
        if source["primary_chapter"] == title:
            score += 20
        ranked.append((score, source["relative_path"], source))
    selected = [source for score, _, source in sorted(ranked, reverse=True) if score > 0]
    return selected[:12]


def source_digest(source):
    title = source["title"]
    rel = source["relative_path"]
    words = source["word_count"]
    pages = source["pages"]
    page_text = f"{pages} paginas, " if pages else ""
    return f"{title}. Fuente original: {rel}. Extension aproximada: {page_text}{words:,} palabras extraidas.".replace(",", ".")


def conceptual_development(chapter, sources):
    source_titles = "; ".join(s["title"] for s in sources[:6]) or "fuentes originales inventariadas"
    focus = chapter["focus"]
    title = chapter["title"]
    paragraphs = [
        f"Este capitulo debe desarrollarse como una unidad de referencia sobre {focus}. La fuente no se trata como una ficha aislada: cada documento se usa para reconstruir definicion, origen, mecanismo, senales, consecuencias, excepciones, herramientas y limites del concepto.",
        f"El principio editorial aplicado aqui es conservar matices. Cuando varias fuentes repiten una idea, la repeticion no se borra sin mas: se distingue entre redundancia superficial y variacion util. La redundancia superficial se fusiona; los ejemplos, contraejemplos, condiciones de aplicacion y lenguaje practico se preservan como subapartados.",
        f"Las fuentes principales que sostienen este bloque en esta primera reconstruccion son: {source_titles}. El siguiente ciclo editorial debe expandir cada fuente con lectura humana o una pasada de extraccion semantica mas larga antes de cerrar una version de cientos de paginas.",
    ]
    if "apego" in title.lower() or "evitacion" in title.lower() or "ansiedad" in title.lower():
        paragraphs.extend([
            "La explicacion central se organiza alrededor del sistema de seguridad. La conducta relacional visible no se interpreta como rasgo moral fijo, sino como estrategia aprendida para recuperar cercania, evitar amenaza, controlar incertidumbre o protegerse de la verguenza.",
            "El manual definitivo debe diferenciar estilo, estado y estrategia: una persona puede tener tendencia ansiosa o evitativa, pero tambien puede activarse de forma situacional segun pareja, contexto, historia y nivel de regulacion disponible.",
        ])
    if "comunicacion" in title.lower() or "conflicto" in title.lower():
        paragraphs.extend([
            "La reparacion se trata como una competencia concreta: reconocer activacion, nombrar la necesidad sin acusar, escuchar impacto, reparar conducta y renegociar el acuerdo. No se reduce a pedir perdon; exige cambio observable.",
            "Los limites se explican como informacion y accion, no como castigo. Un limite sano aclara que conducta se acepta, que consecuencia protectora se aplicara y que parte pertenece a cada persona.",
        ])
    if "masculino" in title.lower() or "polaridad" in title.lower():
        paragraphs.extend([
            "La polaridad se conserva solo cuando puede formularse sin dogma de genero: direccion sin control, receptividad sin pasividad, deseo sin manipulacion, liderazgo sin superioridad y vulnerabilidad sin colapso.",
            "El material masculino queda separado para evitar mezclar consejos especificos con principios generales. Cualquier recomendacion se filtra por consentimiento, reciprocidad, responsabilidad emocional y respeto a la autonomia de la otra persona.",
        ])
    if "ejercicios" in title.lower():
        paragraphs.extend([
            "Se incluyen ejercicios de escritura, discernimiento, comunicacion, observacion de patrones, limites y toma de decisiones. Por instruccion actual, se excluyen meditaciones, breathwork y ejercicios guiados de respiracion.",
            "Cada herramienta debe indicar objetivo, cuando usarla, cuando no usarla, pasos, senales de progreso y riesgos de uso compulsivo o evitativo.",
        ])
    return paragraphs


def add_bullets(story, items, styles):
    if not items:
        return
    flow = [ListItem(Paragraph(clean(item), styles["BulletText"]), leftIndent=10) for item in items]
    story.append(ListFlowable(flow, bulletType="bullet", leftIndent=16, bulletFontSize=7))


def add_page_number(canvas, doc):
    canvas.saveState()
    canvas.setFont("Arial", 8)
    canvas.setFillColor(colors.HexColor("#555555"))
    canvas.drawString(2 * cm, 1.1 * cm, "Manual maestro AAQ - fuentes originales")
    canvas.drawRightString(A4[0] - 2 * cm, 1.1 * cm, str(doc.page))
    canvas.restoreState()


def build_pdf(sources, pdf_path):
    font = "/System/Library/Fonts/Supplemental/Arial.ttf"
    bold = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
    pdfmetrics.registerFont(TTFont("Arial", font))
    pdfmetrics.registerFont(TTFont("Arial-Bold", bold))

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("TitleMain", fontName="Arial-Bold", fontSize=25, leading=30, alignment=TA_CENTER, spaceAfter=18, textColor=colors.HexColor("#1F2933")))
    styles.add(ParagraphStyle("Subtitle", fontName="Arial", fontSize=10.5, leading=15, alignment=TA_CENTER, textColor=colors.HexColor("#52606D")))
    styles.add(ParagraphStyle("Chapter", fontName="Arial-Bold", fontSize=17, leading=21, spaceBefore=10, spaceAfter=10, textColor=colors.HexColor("#102A43")))
    styles.add(ParagraphStyle("H2Local", fontName="Arial-Bold", fontSize=12.2, leading=15, spaceBefore=8, spaceAfter=5, textColor=colors.HexColor("#243B53")))
    styles.add(ParagraphStyle("BodyLocal", fontName="Arial", fontSize=9.2, leading=13.4, alignment=TA_JUSTIFY, spaceAfter=5))
    styles.add(ParagraphStyle("Small", fontName="Arial", fontSize=7.5, leading=10, textColor=colors.HexColor("#52606D")))
    styles.add(ParagraphStyle("BulletText", fontName="Arial", fontSize=8.6, leading=12.2))

    doc = BaseDocTemplate(str(pdf_path), pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm, topMargin=2 * cm, bottomMargin=1.8 * cm)
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="body", frames=[frame], onPage=add_page_number)])

    story = []
    story.append(Spacer(1, 4 * cm))
    story.append(Paragraph("Manual maestro AAQ - fuentes originales", styles["TitleMain"]))
    story.append(Paragraph("Reconstruccion source-only desde /Users/pedro/Desktop/Archivo", styles["Subtitle"]))
    story.append(Paragraph(f"Generado el {datetime.now().strftime('%Y-%m-%d %H:%M')} desde {len(sources)} archivos originales PDF/DOCX.", styles["Subtitle"]))
    story.append(PageBreak())

    story.append(Paragraph("Nota de alcance", styles["Chapter"]))
    scope = [
        "Esta version reemplaza el uso de brain/knowledge como fuente. La extraccion, inventario y trazabilidad parten exclusivamente de los PDF y DOCX originales localizados en /Users/pedro/Desktop/Archivo.",
        "La peticion final exige una obra de referencia exhaustiva, potencialmente de cientos de paginas. Esta entrega deja creada la base source-only: inventario completo, texto extraido por archivo, trazabilidad y una primera arquitectura editorial profunda para continuar sin volver al knowledge brain.",
        "Por derechos de autor y calidad editorial, el manual no vuelca libros completos de terceros palabra por palabra. El siguiente ciclo debe ampliar con sintesis semantica propia, conservando ideas, matices y ejemplos sin copiar masivamente texto protegido.",
        "Regla actual aplicada: incluir ejercicios practicos no meditativos; excluir meditaciones, breathwork y ejercicios guiados de respiracion.",
    ]
    for paragraph in scope:
        story.append(Paragraph(clean(paragraph), styles["BodyLocal"]))

    type_counts = Counter(s["type"] for s in sources)
    total_words = sum(s["word_count"] for s in sources)
    data = [["Metrica", "Valor"], ["Archivos originales", str(len(sources))], ["PDF", str(type_counts.get("pdf", 0))], ["DOCX", str(type_counts.get("docx", 0))], ["Palabras extraidas", f"{total_words:,}".replace(",", ".")]]
    table = Table(data, colWidths=[7 * cm, 6 * cm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E6F0F3")),
        ("FONTNAME", (0, 0), (-1, -1), "Arial"),
        ("FONTNAME", (0, 0), (-1, 0), "Arial-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#BCCCDC")),
        ("PADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(Spacer(1, 8))
    story.append(table)
    story.append(PageBreak())

    story.append(Paragraph("Indice", styles["Chapter"]))
    for i, chapter in enumerate(CHAPTERS, start=1):
        story.append(Paragraph(f"{i}. {clean(chapter['title'])}", styles["BodyLocal"]))
    story.append(Paragraph(f"{len(CHAPTERS) + 1}. Inconsistencias, contradicciones y decisiones editoriales", styles["BodyLocal"]))
    story.append(Paragraph(f"{len(CHAPTERS) + 2}. Mapa completo de fuentes originales", styles["BodyLocal"]))
    story.append(PageBreak())

    for i, chapter in enumerate(CHAPTERS, start=1):
        selected = chapter_sources(sources, chapter)
        story.append(Paragraph(f"{i}. {clean(chapter['title'])}", styles["Chapter"]))
        story.append(Paragraph("Desarrollo editorial", styles["H2Local"]))
        for paragraph in conceptual_development(chapter, selected):
            story.append(Paragraph(clean(paragraph), styles["BodyLocal"]))

        story.append(Paragraph("Preguntas que este capitulo debe responder en la version exhaustiva", styles["H2Local"]))
        add_bullets(story, [
            "Cual es la definicion completa del concepto y que problemas intenta explicar.",
            "Cual es el origen psicologico, relacional o corporal del patron.",
            "Que mecanismos internos lo mantienen y como evoluciona en el tiempo.",
            "Como se identifica en lenguaje, conducta, eleccion de pareja y conflicto.",
            "Que errores de interpretacion aparecen con frecuencia.",
            "Que excepciones o limites tiene esta explicacion.",
            "Que herramientas practicas ayudan y cuando podrian ser contraproducentes.",
        ], styles)

        story.append(Paragraph("Fuentes originales ancla", styles["H2Local"]))
        for source in selected[:10]:
            story.append(Paragraph(clean(source_digest(source)), styles["Small"]))
        story.append(PageBreak())

    story.append(Paragraph(f"{len(CHAPTERS) + 1}. Inconsistencias, contradicciones y decisiones editoriales", styles["Chapter"]))
    notes = [
        ("Meditaciones y breathwork", "La conversacion tuvo dos instrucciones distintas: primero incluirlos, despues excluirlos. Se aplica la instruccion mas reciente: ejercicios si, meditaciones y breathwork no."),
        ("Fuentes originales vs knowledge brain", "La entrega anterior uso la capa curada brain/knowledge. Esta reconstruccion la sustituye por PDF/DOCX originales y conserva los textos extraidos en una carpeta externa de trazabilidad."),
        ("Exhaustividad vs copyright", "El objetivo es que el manual reemplace la consulta de originales para entender el sistema. Eso debe lograrse con redaccion propia y sintesis profunda, no copiando libros o transcripciones completos."),
        ("Polaridad vs seguridad", "Los materiales de masculino/femenino deben filtrarse por consentimiento, reciprocidad y no dogmatismo. La polaridad se conserva como dinamica relacional, no como jerarquia fija."),
    ]
    for title, body in notes:
        story.append(Paragraph(clean(title), styles["H2Local"]))
        story.append(Paragraph(clean(body), styles["BodyLocal"]))
    story.append(PageBreak())

    story.append(Paragraph(f"{len(CHAPTERS) + 2}. Mapa completo de fuentes originales", styles["Chapter"]))
    by_chapter = defaultdict(list)
    for source in sources:
        by_chapter[source["primary_chapter"]].append(source)
    for chapter in CHAPTERS:
        title = chapter["title"]
        story.append(Paragraph(clean(title), styles["H2Local"]))
        for source in sorted(by_chapter.get(title, []), key=lambda s: s["relative_path"]):
            story.append(Paragraph(clean(source_digest(source)), styles["Small"]))

    doc.build(story)


def write_traceability(sources):
    TRACE_OUT.mkdir(parents=True, exist_ok=True)
    SOURCE_OUT.mkdir(parents=True, exist_ok=True)
    csv_path = TRACE_OUT / "original_source_manual_traceability.csv"
    json_path = TRACE_OUT / "original_source_manual_traceability.json"
    with csv_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["relative_path", "type", "pages", "word_count", "primary_chapter", "text_path"])
        writer.writeheader()
        for source in sources:
            writer.writerow({key: source[key] for key in writer.fieldnames})
    summary = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "source_root": str(SOURCE_ROOT),
        "source_file_count": len(sources),
        "type_counts": Counter(s["type"] for s in sources),
        "total_extracted_words": sum(s["word_count"] for s in sources),
        "outputs": {
            "extracted_text_root": str(TEXT_OUT),
            "traceability_csv": str(csv_path),
        },
    }
    json_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    return csv_path, json_path


def main():
    PDF_OUT.mkdir(parents=True, exist_ok=True)
    sources = extract_sources()
    pdf_path = PDF_OUT / "manual_maestro_fuentes_originales_aaq.pdf"
    build_pdf(sources, pdf_path)
    csv_path, json_path = write_traceability(sources)
    print(pdf_path)
    print(csv_path)
    print(json_path)
    print(TEXT_OUT)


if __name__ == "__main__":
    main()
